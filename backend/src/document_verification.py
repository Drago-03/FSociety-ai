import os
import uuid
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from fastapi import UploadFile, HTTPException
import aiohttp
import PyPDF2
import docx
import magic
import hashlib
from io import BytesIO

from model_utils import ContentModerator
from cloud_storage import CloudStorage

logger = logging.getLogger(__name__)

class DocumentVerifier:
    def __init__(self):
        self.content_moderator = ContentModerator()
        self.cloud_storage = CloudStorage()
        self.mime = magic.Magic(mime=True)
        self.trusted_sources = [
            "https://www.govinfo.gov",
            "https://www.irs.gov",
            "https://www.sec.gov",
            "https://www.justice.gov",
            "https://www.ftc.gov"
        ]
    
    async def verify_document(self, file: UploadFile) -> Dict[str, Any]:
        """Verify a document for authenticity and detect potential issues"""
        try:
            # Read file content
            content = await file.read()
            file_size = len(content)
            file_hash = hashlib.sha256(content).hexdigest()
            
            # Detect file type
            file_type = self.mime.from_buffer(content)
            
            # Extract metadata and text
            metadata = self.extract_metadata(file.filename, content, file_type)
            text_content = self.extract_text(content, file_type)
            
            # Store file in cloud storage
            file_id = str(uuid.uuid4())
            storage_path = f"documents/{file_id}/{file.filename}"
            await self.cloud_storage.upload_file(storage_path, BytesIO(content))
            
            # Analyze content for issues
            content_analysis = self.content_moderator.analyze_toxicity(text_content)
            similarity_results = await self.check_against_trusted_sources(text_content)
            
            # Determine authenticity based on analysis
            issues = self.identify_issues(metadata, content_analysis, similarity_results)
            is_authentic = len(issues) == 0
            confidence = self.calculate_confidence(metadata, content_analysis, similarity_results)
            
            return {
                "document_id": file_id,
                "filename": file.filename,
                "file_size": file_size,
                "file_type": file_type,
                "file_hash": file_hash,
                "upload_timestamp": datetime.utcnow().isoformat(),
                "is_authentic": is_authentic,
                "confidence": confidence,
                "category": self.detect_document_category(file.filename, text_content),
                "issues": issues,
                "metadata": metadata,
                "content_analysis": content_analysis,
                "storage_path": storage_path
            }
        except Exception as e:
            logger.error(f"Error verifying document: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Document verification failed: {str(e)}")
    
    def extract_metadata(self, filename: str, content: bytes, file_type: str) -> Dict[str, Any]:
        """Extract metadata from document content"""
        metadata = {
            "filename": filename,
            "file_size": len(content),
            "file_type": file_type,
            "extraction_time": datetime.utcnow().isoformat()
        }
        
        try:
            if "pdf" in file_type.lower():
                # Extract PDF metadata
                pdf = PyPDF2.PdfReader(BytesIO(content))
                metadata.update({
                    "page_count": len(pdf.pages),
                    "pdf_version": pdf.pdf_header,
                    "encrypted": pdf.is_encrypted
                })
                
                if pdf.metadata:
                    for key, value in pdf.metadata.items():
                        if key.startswith('/'):
                            clean_key = key[1:].lower()
                            metadata[clean_key] = str(value)
            
            elif "word" in file_type.lower() or "docx" in file_type.lower():
                # Extract DOCX metadata
                doc = docx.Document(BytesIO(content))
                core_props = doc.core_properties
                metadata.update({
                    "author": core_props.author,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision,
                    "title": core_props.title,
                    "paragraph_count": len(doc.paragraphs)
                })
        except Exception as e:
            logger.warning(f"Error extracting detailed metadata: {str(e)}")
        
        return metadata
    
    def extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text content from document"""
        try:
            if "pdf" in file_type.lower():
                # Extract text from PDF
                pdf = PyPDF2.PdfReader(BytesIO(content))
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
                return text
            
            elif "word" in file_type.lower() or "docx" in file_type.lower():
                # Extract text from DOCX
                doc = docx.Document(BytesIO(content))
                return "\n".join([para.text for para in doc.paragraphs])
            
            elif "text" in file_type.lower():
                # Plain text
                return content.decode('utf-8', errors='replace')
            
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return ""
    
    async def check_against_trusted_sources(self, text: str) -> Dict[str, Any]:
        """Check document content against trusted sources"""
        # In a real implementation, this would query trusted sources
        # or use a verification service to check content authenticity
        
        # For demonstration, we'll implement a simplified version
        try:
            # Extract key phrases for verification
            key_phrases = self.extract_key_phrases(text)
            
            results = {
                "verified_phrases": [],
                "unverified_phrases": [],
                "matched_sources": [],
                "overall_match_score": 0.0
            }
            
            # In a real implementation, we would check these phrases against trusted sources
            # For now, we'll simulate the results
            verified_count = 0
            for phrase in key_phrases:
                if len(phrase) > 10 and hash(phrase) % 3 != 0:  # Simple simulation
                    results["verified_phrases"].append(phrase)
                    verified_count += 1
                else:
                    results["unverified_phrases"].append(phrase)
            
            # Calculate match score
            if key_phrases:
                results["overall_match_score"] = verified_count / len(key_phrases)
            
            # Simulate matched sources
            if results["verified_phrases"]:
                results["matched_sources"] = [source for source in self.trusted_sources 
                                           if hash(source + text) % 4 != 0]  # Simple simulation
            
            return results
        except Exception as e:
            logger.error(f"Error checking against trusted sources: {str(e)}")
            return {
                "verified_phrases": [],
                "unverified_phrases": [],
                "matched_sources": [],
                "overall_match_score": 0.0,
                "error": str(e)
            }
    
    def extract_key_phrases(self, text: str) -> List[str]:
        """Extract key phrases from text for verification"""
        # In a real implementation, this would use NLP techniques
        # to extract important phrases for verification
        
        # Simple implementation: split by sentences and paragraphs
        phrases = []
        paragraphs = text.split('\n')
        for para in paragraphs:
            if len(para.strip()) > 20:  # Ignore very short paragraphs
                sentences = para.split('.')
                for sentence in sentences:
                    if len(sentence.strip()) > 30:  # Only consider substantial sentences
                        phrases.append(sentence.strip())
        
        # Limit to a reasonable number of phrases
        return phrases[:20]
    
    def identify_issues(self, metadata: Dict[str, Any], 
                       content_analysis: Dict[str, Any],
                       similarity_results: Dict[str, Any]) -> List[str]:
        """Identify potential issues in the document"""
        issues = []
        
        # Check metadata for issues
        if metadata.get("encrypted", False):
            issues.append("Document is encrypted which may hide content")
        
        # Check for suspicious modification patterns
        if metadata.get("created") and metadata.get("modified"):
            created = datetime.fromisoformat(metadata["created"])
            modified = datetime.fromisoformat(metadata["modified"])
            if (modified - created).total_seconds() < 60:  # Less than a minute difference
                issues.append("Suspicious modification pattern detected")
        
        # Check content analysis
        if content_analysis.get("toxicity", 0) > 0.7:
            issues.append("Document contains potentially harmful content")
        
        # Check similarity results
        if similarity_results.get("overall_match_score", 0) < 0.3 and similarity_results.get("verified_phrases"):
            issues.append("Low match with trusted sources")
        
        if len(similarity_results.get("unverified_phrases", [])) > len(similarity_results.get("verified_phrases", [])):
            issues.append("Multiple unverified statements detected")
        
        return issues
    
    def calculate_confidence(self, metadata: Dict[str, Any],
                           content_analysis: Dict[str, Any],
                           similarity_results: Dict[str, Any]) -> float:
        """Calculate confidence score for document authenticity"""
        # Base confidence score
        confidence = 0.5
        
        # Adjust based on metadata
        if metadata.get("author") and metadata.get("created") and metadata.get("title"):
            confidence += 0.1  # Complete metadata is good
        
        # Adjust based on content analysis
        toxicity = content_analysis.get("toxicity", 0)
        confidence -= toxicity * 0.2  # Higher toxicity reduces confidence
        
        # Adjust based on similarity results
        match_score = similarity_results.get("overall_match_score", 0)
        confidence += match_score * 0.3  # Higher match with trusted sources increases confidence
        
        # Ensure confidence is between 0 and 1
        return max(0.0, min(1.0, confidence))
    
    def detect_document_category(self, filename: str, text_content: str) -> str:
        """Detect the category of a document based on filename and content"""
        filename_lower = filename.lower()
        text_lower = text_content.lower()
        
        # Check filename first
        if any(term in filename_lower for term in ["contract", "agreement", "terms", "legal"]):
            return "legal_document"
        elif any(term in filename_lower for term in ["report", "financial", "balance", "income", "statement"]):
            return "financial_document"
        elif any(term in filename_lower for term in ["policy", "handbook", "manual", "guide"]):
            return "policy_document"
        elif any(term in filename_lower for term in ["id", "passport", "license", "certificate"]):
            return "identity_document"
        
        # Check content if filename doesn't provide enough information
        if any(term in text_lower for term in ["contract", "agreement", "parties", "hereby", "terms", "conditions"]):
            return "legal_document"
        elif any(term in text_lower for term in ["financial", "revenue", "profit", "loss", "balance", "income", "statement"]):
            return "financial_document"
        elif any(term in text_lower for term in ["policy", "procedure", "guideline", "handbook"]):
            return "policy_document"
        elif any(term in text_lower for term in ["identification", "passport", "license", "certificate", "birth"]):
            return "identity_document"
        
        # Default category
        return "general_document"


class WebScraper:
    def __init__(self):
        self.session = None
        self.headers = {
            "User-Agent": "FSociety-AI-Verification-Service/1.0",
            "Accept": "text/html,application/xhtml+xml,application/xml",
            "Accept-Language": "en-US,en;q=0.9"
        }
    
    async def initialize(self):
        """Initialize aiohttp session"""
        if self.session is None or self.session.closed:
            self.session = aiohttp.ClientSession(headers=self.headers)
    
    async def close(self):
        """Close aiohttp session"""
        if self.session and not self.session.closed:
            await self.session.close()
    
    async def scrape_url(self, url: str) -> Dict[str, Any]:
        """Scrape content from a URL"""
        await self.initialize()
        
        try:
            async with self.session.get(url, timeout=30) as response:
                if response.status != 200:
                    return {
                        "success": False,
                        "status_code": response.status,
                        "error": f"Failed to fetch URL: HTTP {response.status}"
                    }
                
                content_type = response.headers.get('Content-Type', '')
                if 'text/html' in content_type:
                    html_content = await response.text()
                    return {
                        "success": True,
                        "url": url,
                        "content_type": content_type,
                        "html_content": html_content,
                        "headers": dict(response.headers),
                        "timestamp": datetime.utcnow().isoformat()
                    }
                elif 'application/pdf' in content_type:
                    pdf_content = await response.read()
                    # Extract text from PDF
                    pdf_text = self.extract_text_from_pdf(pdf_content)
                    return {
                        "success": True,
                        "url": url,
                        "content_type": content_type,
                        "text_content": pdf_text,
                        "headers": dict(response.headers),
                        "timestamp": datetime.utcnow().isoformat()
                    }
                else:
                    # For other content types, just return basic info
                    return {
                        "success": True,
                        "url": url,
                        "content_type": content_type,
                        "headers": dict(response.headers),
                        "timestamp": datetime.utcnow().isoformat(),
                        "note": "Content type not supported for detailed extraction"
                    }
        except Exception as e:
            logger.error(f"Error scraping URL {url}: {str(e)}")
            return {
                "success": False,
                "url": url,
                "error": str(e),
                "timestamp": datetime.utcnow().isoformat()
            }
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf = PyPDF2.PdfReader(BytesIO(pdf_content))
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    async def scrape_multiple_urls(self, urls: List[str]) -> List[Dict[str, Any]]:
        """Scrape content from multiple URLs in parallel"""
        await self.initialize()
        
        tasks = [self.scrape_url(url) for url in urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({
                    "success": False,
                    "url": urls[i],
                    "error": str(result),
                    "timestamp": datetime.utcnow().isoformat()
                })
            else:
                processed_results.append(result)
        
        return processed_results